import os
import io
import uuid
from flask import Flask, render_template, request, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from html.parser import HTMLParser
import platform

class DocxRTEParser(HTMLParser):
    def __init__(self, paragraph):
        super().__init__()
        self.p = paragraph
        self.bold = False
        self.italic = False
        self.underline = False
        self.highlight = False
        self.highlight_color = WD_COLOR_INDEX.YELLOW
        self.has_content = False

    def handle_starttag(self, tag, attrs):
        if tag in ('b', 'strong'): self.bold = True
        elif tag in ('i', 'em'): self.italic = True
        elif tag == 'u': self.underline = True
        elif tag in ('div', 'p', 'br'):
            if self.has_content:
                # Add a space between lines if user pressed enter
                run = self.p.add_run(' ')
                self.has_content = False
        elif tag in ('span', 'mark', 'font'):
            for attr, val in attrs:
                if attr == 'style' and 'background' in val.lower():
                    self.highlight = True
                    val_lower = val.lower()
                    if 'yellow' in val_lower or '255, 255, 0' in val_lower or '#ffff00' in val_lower:
                        self.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif 'lime' in val_lower or 'green' in val_lower or '0, 255, 0' in val_lower or '#00ff00' in val_lower:
                        self.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
                    elif 'cyan' in val_lower or '0, 255, 255' in val_lower or '#00ffff' in val_lower or 'turquoise' in val_lower:
                        self.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    elif 'pink' in val_lower or '255, 192, 203' in val_lower or '#ffc0cb' in val_lower:
                        self.highlight_color = WD_COLOR_INDEX.PINK
                    else:
                        self.highlight_color = WD_COLOR_INDEX.YELLOW
                elif attr == 'class' and 'highlight' in val.lower():
                    self.highlight = True
                    self.highlight_color = WD_COLOR_INDEX.YELLOW

    def handle_endtag(self, tag):
        if tag in ('b', 'strong'): self.bold = False
        elif tag in ('i', 'em'): self.italic = False
        elif tag == 'u': self.underline = False
        elif tag in ('span', 'mark', 'font'):
            self.highlight = False
            self.highlight_color = WD_COLOR_INDEX.YELLOW

    def handle_data(self, data):
        if data:
            self.has_content = True
            run = self.p.add_run(data)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = self.bold
            run.italic = self.italic
            run.underline = self.underline
            if self.highlight:
                run.font.highlight_color = getattr(self, 'highlight_color', WD_COLOR_INDEX.YELLOW)

import subprocess

def convert(input_path, output_path):
    if platform.system() == 'Windows':
        import pythoncom
        try:
            from docx2pdf import convert as windows_convert
            pythoncom.CoInitialize()
            windows_convert(input_path, output_path)
            pythoncom.CoUninitialize()
        except ImportError:
            raise Exception("docx2pdf is not installed on Windows")
    else:
        # Linux / Render approach using LibreOffice
        try:
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(output_path), input_path], check=True, capture_output=True)
            # libreoffice saves it exactly as input_name.pdf, rename it to what we actually wanted
            original_output = os.path.join(os.path.dirname(output_path), os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
            if original_output != output_path and os.path.exists(original_output):
                os.rename(original_output, output_path)
        except Exception as e:
            raise Exception(f"LibreOffice conversion failed: {str(e)}")

app = Flask(__name__)
app.secret_key = 'super_secret_key'
app.config['UPLOAD_FOLDER'] = os.path.join(app.root_path, 'uploads')
app.config['TEMP_FOLDER'] = os.path.join(app.root_path, 'temp')
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['TEMP_FOLDER'], exist_ok=True)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generate_report', methods=['POST'])
def generate_report():
    # Helper to clean strings
    def clean(s):
        return s.strip() if s else ""

    # Parse request form data
    report_title = clean(request.form.get('reportTitle', 'REPORT')).upper()
    event_title = clean(request.form.get('eventTitle', ''))
    event_date = clean(request.form.get('eventDate', ''))
    venue = clean(request.form.get('venue', ''))
    event_coordinator = clean(request.form.get('eventCoordinator', ''))
    faculty_coordinators = [f.strip() for f in request.form.getlist('facultyCoordinators') if f.strip()]
    resource_person = clean(request.form.get('resourcePersonName', ''))
    designation = clean(request.form.get('designation', ''))
    objectives = [o.strip() for o in request.form.getlist('objectives') if o.strip()]
    target_audience = clean(request.form.get('targetAudience', ''))
    num_participants = clean(request.form.get('numParticipants', ''))
    outcomes = [o.strip() for o in request.form.getlist('outcomes') if o.strip()]
    conclusion = clean(request.form.get('conclusionParagraph', ''))

    # Event mapping with course Objectives: text + level (Low/Medium/High) per row
    co_texts  = request.form.getlist('courseObjective')
    co_levels = [request.form.get(f'courseObjLevel_{i}', '') for i in range(1, 6)]
    course_objectives = [
        (t.strip(), l.strip())
        for t, l in zip(co_texts, co_levels)
        if t.strip()
    ]

    # Theory Outcomes: text + Yes/No per row
    to_texts  = request.form.getlist('theoryOutcome')
    to_levels = [request.form.get(f'theoryOutcomeYN_{i}', '') for i in range(1, 5)]
    theory_outcomes = [
        (t.strip(), l.strip())
        for t, l in zip(to_texts, to_levels)
        if t.strip()
    ]

    # Lab Outcomes (Optional): text + Yes/No per row
    lo_texts  = request.form.getlist('labOutcome')
    lo_levels = [request.form.get(f'labOutcomeYN_{i}', '') for i in range(1, 5)]
    lab_outcomes = [
        (t.strip(), l.strip())
        for t, l in zip(lo_texts, lo_levels)
        if t.strip()
    ]
    hod_name = clean(request.form.get('hodName', ''))

    doc = Document()
    
    # ---- Page margins ----
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # ---------------- FOOTER ----------------
    footer = section.footer
    p_foot = footer.paragraphs[0]
    p_foot.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    
    run_foot = p_foot.add_run("CSE(DS)\twww.rgmcet.edu.in")
    run_foot.font.name = 'Arial'
    run_foot.font.size = Pt(11)
    run_foot.font.color.rgb = RGBColor(120, 120, 120)
    
    # Add top border to footer
    pPr = p_foot._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '12')
    top.set(qn('w:color'), 'auto')
    pBdr.append(top)
    pPr.append(pBdr)

    BLUE = RGBColor(46, 116, 181)
    logo_path = os.path.join(app.root_path, 'static', 'images', 'logo_rgm.png')

    # ============== RENDER HEADER AS IMAGE (matches web HTML header exactly) ==============
    from PIL import Image, ImageDraw, ImageFont

    def create_header_image():
        """Render the exact HTML header as a high-quality PNG image."""
        # 3x scale for crisp print quality (effective ~300 DPI)
        SCALE = 3
        # Usable page width = 6.5 inches at 96 DPI
        canvas_w = int(6.5 * 96 * SCALE)   # ~1872 px
        canvas_h = int(1.4 * 96 * SCALE)   # ~403 px
        
        img = Image.new('RGBA', (canvas_w, canvas_h), (255, 255, 255, 255))
        draw = ImageDraw.Draw(img)
        
        # Load fonts (Georgia from Windows)
        font_dir = r'C:\Windows\Fonts'
        try:
            # Match CSS: h2 = 1.4rem (~16pt), p = 1.1rem (~12pt), h3 = 1.25rem (~14pt)
            font_college = ImageFont.truetype(os.path.join(font_dir, 'georgia.ttf'), int(16 * SCALE))
            font_sub = ImageFont.truetype(os.path.join(font_dir, 'georgia.ttf'), int(12 * SCALE))
            font_dept = ImageFont.truetype(os.path.join(font_dir, 'georgiab.ttf'), int(13 * SCALE))
        except Exception:
            font_college = ImageFont.load_default()
            font_sub = font_college
            font_dept = font_college
        
        # Colors matching CSS
        gray = (119, 119, 119)   # #777
        dark = (102, 102, 102)   # #666
        
        # ---- Logo on the left ----
        logo_area_w = int(1.1 * 96 * SCALE)  # ~1.1 inches for logo column
        if os.path.exists(logo_path):
            logo_img = Image.open(logo_path).convert('RGBA')
            max_logo_h = canvas_h - int(20 * SCALE)
            max_logo_w = logo_area_w - int(15 * SCALE)
            ratio = min(max_logo_w / logo_img.width, max_logo_h / logo_img.height)
            new_w = int(logo_img.width * ratio)
            new_h = int(logo_img.height * ratio)
            logo_resized = logo_img.resize((new_w, new_h), Image.LANCZOS)
            # Center logo vertically, align left with small padding
            lx = (logo_area_w - new_w) // 2
            ly = (canvas_h - new_h) // 2
            img.paste(logo_resized, (lx, ly), logo_resized)
        
        # ---- Text lines on the right, centered ----
        text_left = logo_area_w + int(10 * SCALE)
        text_right = canvas_w - int(5 * SCALE)
        text_center_x = (text_left + text_right) // 2
        
        lines = [
            ("Rajeev Gandhi Memorial College of Engineering & Technology", font_college, gray),
            ("(Autonomous)", font_sub, gray),
            ("Nandyal-518501, Nandyal \u2013 Dt. AP. India", font_sub, gray),
            ("Department of Computer Science & Engineering (Data Science)", font_dept, dark),
        ]
        
        # Measure text heights
        line_metrics = []
        for text, font, color in lines:
            bbox = draw.textbbox((0, 0), text, font=font)
            tw = bbox[2] - bbox[0]
            th = bbox[3] - bbox[1]
            line_metrics.append((tw, th))
        
        line_gap = int(7 * SCALE)
        # Extra gap before department line (matching CSS margin: 8px 0)
        dept_extra_gap = int(5 * SCALE)
        total_h = sum(m[1] for m in line_metrics) + line_gap * (len(lines) - 2) + line_gap + dept_extra_gap
        
        y = (canvas_h - total_h) // 2
        
        for i, (text, font, color) in enumerate(lines):
            tw, th = line_metrics[i]
            x = text_center_x - tw // 2
            draw.text((x, y), text, fill=color, font=font)
            y += th + line_gap
            # Add extra gap before the department line
            if i == 2:
                y += dept_extra_gap
        
        # Convert to RGB for saving
        result = Image.new('RGB', img.size, (255, 255, 255))
        result.paste(img, mask=img.split()[3])
        return result

    # Create and save header image
    header_img = create_header_image()
    header_img_path = os.path.join(app.config['TEMP_FOLDER'], f"header_{uuid.uuid4().hex[:8]}.png")
    header_img.save(header_img_path, 'PNG', dpi=(300, 300))
    
    # ============== PUT HEADER IN WORD HEADER SECTION (repeats on ALL pages) ==============
    header = section.header
    header.is_linked_to_previous = False
    
    # Clear any default paragraphs
    for p in header.paragraphs:
        p.clear()
    
    # Insert header image in the Word header
    p_header = header.paragraphs[0]
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.space_before = Pt(0)
    p_header.paragraph_format.space_after = Pt(0)
    run_h = p_header.add_run()
    run_h.add_picture(header_img_path, width=Inches(6.5))
    
    # Clean up temp header image
    if os.path.exists(header_img_path):
        os.remove(header_img_path)

    # Blue horizontal line — 0 line gap before, 1 line gap after header image
    p_line = header.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(12)
    pBdr2 = OxmlElement('w:pBdr')
    bottom2 = OxmlElement('w:bottom')
    bottom2.set(qn('w:val'), 'single')
    bottom2.set(qn('w:sz'), '14')
    bottom2.set(qn('w:space'), '1')
    bottom2.set(qn('w:color'), '2E74B5')
    pBdr2.append(bottom2)
    p_line._p.get_or_add_pPr().append(pBdr2)
    # ---------------- TITLE (1 line gap after header/blue line) ----------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(report_title)
    r_title.font.name = 'Times New Roman'
    r_title.bold = True
    r_title.underline = True
    r_title.font.size = Pt(18)
    r_title.font.color.rgb = BLUE

    # ---------------- BODY ----------------
    def add_section_header(num, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(10)
        r = p.add_run(f"{num}. {text}")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(16)
        r.font.color.rgb = BLUE

    # 1. Title of the Event : "event_title" (inline on same line)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(10)
    r_lbl = p.add_run("1. Title of the Event : ")
    r_lbl.font.name = 'Times New Roman'
    r_lbl.font.size = Pt(16)
    r_lbl.font.color.rgb = BLUE
    r_val = p.add_run(event_title)
    r_val.font.name = 'Times New Roman'
    r_val.font.size = Pt(14)
    r_val.bold = True
    r_val.underline = True

    # 2. Date and Venue
    add_section_header(2, "Date and Venue")
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run("Date: ")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = p.add_run(event_date)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.bold = True
    r2.underline = True

    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run("Venue: ")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = p.add_run(venue)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.bold = True
    r2.underline = True

    # 3. Coordinator(s)
    add_section_header(3, "Coordinator(s)")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    r1 = p.add_run("   Event Coordinator: ")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = p.add_run(event_coordinator)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.bold = True
    r2.underline = True

    # Only print Faculty Coordinator(s) if names were provided
    if faculty_coordinators:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)

        r_lbl = p.add_run("•  Faculty Coordinator(s):")
        r_lbl.font.name = 'Times New Roman'
        r_lbl.font.size = Pt(12)
        r_lbl.bold = True

        for fc in faculty_coordinators:
            p.add_run().add_break()
            p.add_run("\t") # One tab space for the names
            r2 = p.add_run(fc)
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)
            r2.bold = True
            r2.underline = True

    # 4. Resource Person(s):
    add_section_header(4, "Resource Person(s):")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    r1 = p.add_run("Name: ")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = p.add_run(resource_person)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    r1 = p.add_run("Designation & Organization: ")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r1.bold = True
    r2 = p.add_run(designation)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

    # 5. Objectives of the Event
    add_section_header(5, "Objectives of the Event")
    for obj in objectives:
        p = doc.add_paragraph()
        r = p.add_run("❑  " + obj)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = True

    # 6. Target Audience
    add_section_header(6, "Target Audience")
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(f"{target_audience} - {num_participants}")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True

    # 7. Outcomes of the Event
    add_section_header(7, "Outcomes of the Event")
    p = doc.add_paragraph()
    r = p.add_run("Students able:")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = True
    for out in outcomes:
        p = doc.add_paragraph()
        r = p.add_run("❑  " + out)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)

    # Dynamic section numbering from 8 onwards
    sec_num = 8

    # 8. Event mapping with course Objectives
    if course_objectives:
        add_section_header(sec_num, "Event mapping with course Objectives")
        sec_num += 1
        for text, level in course_objectives:
            p = doc.add_paragraph()
            r1 = p.add_run(f"❑  {text} - ")
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(12)
            r1.bold = False
            r2 = p.add_run(level)
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)
            r2.bold = True

    # 9. Theory Outcomes
    if theory_outcomes:
        add_section_header(sec_num, "Theory Outcomes")
        sec_num += 1
        for text, yn in theory_outcomes:
            p = doc.add_paragraph()
            r1 = p.add_run(f"❑  {text} - ")
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(12)
            r1.bold = False
            r2 = p.add_run(yn)
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)
            r2.bold = True

    # 10. Lab Outcomes (Optional — before Conclusion)
    if lab_outcomes:
        add_section_header(sec_num, "Lab Outcomes")
        sec_num += 1
        for text, yn in lab_outcomes:
            p = doc.add_paragraph()
            r1 = p.add_run(f"❑  {text} - ")
            r1.font.name = 'Times New Roman'
            r1.font.size = Pt(12)
            r1.bold = False
            r2 = p.add_run(yn)
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)
            r2.bold = True

    # Conclusion — always the next available number
    if conclusion:
        add_section_header(sec_num, "Conclusion")
        sec_num += 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        
        # Parse the HTML conclusion sent from textarea editor to retain formatting
        parser = DocxRTEParser(p)
        parser.feed(conclusion)
        
    # ---------------- PHOTOGRAPHS (Mostly Page 3) ----------------
    if any(request.files.get(f'photo{i}') for i in range(1, 5)):
        # Force photographs to a new page completely for professional aesthetic
        doc.add_page_break()
        
        # Pushing the section 1 line down
        doc.add_paragraph()

        add_section_header(sec_num, "Photographs")
        sec_num += 1
        
        # Create a professional 2x2 table grid for images
        photo_table = doc.add_table(rows=2, cols=2)
        photo_table.style = 'Table Grid'
        photo_table.alignment = WD_TAB_ALIGNMENT.CENTER
        
        # Optimize cell properties
        for row in photo_table.rows:
            # Set minimum row height directly in XML if needed, but python-docx autoscales heights well.
            for cell in row.cells:
                cell.width = Inches(3.25)
                # Add small padding to prevent images from touching borders
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(5)
                cp.paragraph_format.space_after = Pt(5)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        photo_inputs = ['photo1', 'photo2', 'photo3', 'photo4']
        for idx, photo_name in enumerate(photo_inputs):
            photo = request.files.get(photo_name)
            if photo and photo.filename:
                r_idx = idx // 2
                c_idx = idx % 2
                
                ext = os.path.splitext(photo.filename)[1]
                if not ext: ext = '.jpg'
                temp_path = os.path.join(app.config['TEMP_FOLDER'], f"img_{uuid.uuid4().hex[:8]}{ext}")
                photo.save(temp_path)
                
                cell = photo_table.cell(r_idx, c_idx)
                cp = cell.paragraphs[0]
                
                try:
                    # Scale down the image slightly to give a clear gap between the grid lines
                    cp.add_run().add_picture(temp_path, width=Inches(3.0))
                except Exception:
                    pass
                finally:
                    # Clean up immediately
                    if os.path.exists(temp_path):
                        os.remove(temp_path)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # ---------------- SIGNATURES ----------------
    num_cols = 3 if faculty_coordinators else 2
    sig_table = doc.add_table(rows=1, cols=num_cols)

    cell_ec  = sig_table.cell(0, 0)
    cell_hod = sig_table.cell(0, num_cols - 1)

    # Event Coordinator
    p_ec = cell_ec.paragraphs[0]
    p_ec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_ec.add_run("\n\n\nEvent coordinator\n")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r2 = p_ec.add_run(f"(Dr. {event_coordinator})" if 'Dr' not in event_coordinator else f"({event_coordinator})")
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)

    # Faculty-Coordinator — only when names are provided
    if faculty_coordinators:
        cell_fc = sig_table.cell(0, 1)
        p_fc = cell_fc.paragraphs[0]
        p_fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_fc.add_run("\n\n\nFaculty-Coordinator\n")
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        for fc in faculty_coordinators:
            r2 = p_fc.add_run(f"{fc}\n")
            r2.bold = True
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(12)

    # HoD
    p_hod = cell_hod.paragraphs[0]
    p_hod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_hod.add_run("\n\n\nHoD\n")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r2 = p_hod.add_run(f"({hod_name})")
    r2.bold = True
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)



    # Clean up title for file name
    safe_title = "".join(c for c in report_title if c.isalnum() or c in (' ', '_', '-')).strip()
    
    # Save the doc to temp storage
    internal_id = str(uuid.uuid4())
    docx_path = os.path.join(app.config['TEMP_FOLDER'], f"{internal_id}.docx")
    doc.save(docx_path)
    
    return jsonify({
        'success': True, 
        'report_id': internal_id, 
        'filename': safe_title if safe_title else "Report"
    })


@app.route('/download/<report_id>/<format>')
def download(report_id, format):
    docx_path = os.path.join(app.config['TEMP_FOLDER'], f"{report_id}.docx")
    if not os.path.exists(docx_path):
        flash("File not found or expired.", "danger")
        return redirect(url_for('index'))
        
    filename = request.args.get('filename', 'Report')
    preview = request.args.get('preview', '0') == '1'
    as_attachment = not preview
    
    if format == 'pdf':
        pdf_path = os.path.join(app.config['TEMP_FOLDER'], f"{report_id}.pdf")
        if not os.path.exists(pdf_path):
            try:
                convert(docx_path, pdf_path)
            except Exception as e:
                return f"PDF Conversion Failed: {str(e)}", 500
                
        return send_file(pdf_path, as_attachment=as_attachment, download_name=f"{filename}.pdf", mimetype='application/pdf')
    else:
        return send_file(docx_path, as_attachment=as_attachment, download_name=f"{filename}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True, port=5000)
